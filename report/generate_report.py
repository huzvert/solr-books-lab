"""Generate the Word lab report for Lab 13.

Produces report/Lab13_Solr_Report.docx with:
  problem statement, dataset description, configuration,
  implementation steps, screenshots (if available),
  observations, challenges, conclusion.

Usage:  python report/generate_report.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "screenshots"
OUT = ROOT / "report" / "Lab13_Solr_Report.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def add_image_if_exists(doc, name, caption):
    f = SCREENSHOTS / name
    if f.exists():
        doc.add_picture(str(f), width=Inches(6))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)
    else:
        add_para(doc, f"[Screenshot placeholder: {name} — {caption}]", italic=True)


def main():
    doc = Document()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Lab 13 — Indexing, Importing and Searching Data in Apache Solr")
    r.bold = True
    r.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Course: CS-347   |   Class: BSCS-13AB   |   Instructor: Dr. Khurram Shahzad\n"
                "Date: 8 May 2026").italic = True

    stu = doc.add_paragraph()
    stu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = stu.add_run("Student: Huzaifa Ali Satti   |   CMS ID: 468629   |   NUST SEECS")
    sr.bold = True

    doc.add_paragraph()

    # 1. Problem statement
    add_heading(doc, "1. Problem Statement")
    add_para(doc,
        "The lab task is to set up Apache Solr, index a real dataset, run a "
        "range of search queries against it, and put a web UI on top. Apache "
        "Solr is a Lucene-based search engine. It speaks HTTP, so most of the "
        "work in the UI is just translating form fields into query parameters. "
        "I picked a real Amazon product catalog as the corpus because it gives "
        "every UI feature (price-range filtering, in-stock facet, brand "
        "drilldown, rating sort) something real to operate on. This report "
        "covers both halves, with enough configuration detail that the grader "
        "can reproduce the setup on a fresh machine."
    )

    # 2. Dataset description
    add_heading(doc, "2. Dataset Description")
    add_para(doc,
        "The dataset is the Bright Data Amazon-products sample "
        "(github.com/luminati-io/Amazon-dataset-samples), a public CSV of "
        "1,000 real Amazon listings with 55 columns each. data/transform_amazon.py "
        "extracts the 11 fields used by this lab and drops 4 rows that are "
        "missing brand, price, or rating. The result is 996 products with all "
        "fields read directly from the source — no hash-derived or synthetic "
        "values.")
    add_para(doc, "Each record is a single product with these fields:")
    fields_table = [
        ("id", "string", "Amazon ASIN (B09NQJFRW6, B0074TRKFI, ...)"),
        ("title", "text_general", "Product name (full-text indexed)"),
        ("brand", "string", "Brand name — used as facet, untokenized so 'New Balance' stays one bucket"),
        ("brand_text", "text_general", "Tokenized copy of brand for full-text search; populated via copyField"),
        ("category", "string", "Top-level Amazon category — facet field"),
        ("subcategory", "string", "Second-level Amazon category — facet field"),
        ("year", "pint", "First-available year, parsed from date_first_available"),
        ("num_reviews", "pint", "Review count from Amazon, sortable"),
        ("price", "pfloat", "Real Amazon price in USD"),
        ("rating", "pfloat", "Real Amazon star rating (1.0-5.0)"),
        ("in_stock", "boolean", "Mapped from the availability column (true unless 'out of stock' / 'unavailable')"),
        ("description", "text_general", "Real product description, truncated to 1500 chars"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Solr type"
    hdr[2].text = "Purpose"
    for f, t_, p in fields_table:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = t_
        row[2].text = p

    add_para(doc, "")
    add_para(doc,
        "Source: github.com/luminati-io/Amazon-dataset-samples. Format: CSV, "
        "996 rows after cleaning, ~2 MB. Top categories are Clothing/Shoes/"
        "Jewelry (216), Home & Kitchen (150), Tools & Home Improvement (112), "
        "Sports & Outdoors (78) and Electronics (75). Brands range from "
        "household names (Skechers, New Balance, BOSCH, adidas) to long-tail "
        "third-party sellers."
    )

    # 3. Configuration details
    add_heading(doc, "3. Configuration Details")
    add_para(doc, "Software stack:", bold=True)
    add_para(doc,
        "* Apache Solr 9.6.1 (SolrCloud mode, 2 nodes, embedded ZooKeeper)\n"
        "* OpenJDK 21\n"
        "* Python 3.13 with Flask 3.x and requests\n"
        "* Windows 11"
    )
    add_para(doc, "Cluster topology:", bold=True)
    add_para(doc,
        "I started Solr with 'solr.cmd -e cloud -noprompt', which provisions a "
        "two-node cluster: node1 on port 8983 and node2 on port 7574, with an "
        "embedded ZooKeeper instance on port 9983 coordinating cluster state. "
        "The 'products' collection is created with numShards=2 and "
        "replicationFactor=2, giving 4 cores total: each shard has a leader "
        "on one node and a replica on the other. With ~996 documents that "
        "puts roughly 498 in each shard. The Flask UI talks to a single "
        "endpoint (http://localhost:8983/solr/products/select) and Solr "
        "transparently fan-outs the query to both shards and merges the "
        "results.")
    add_para(doc, "Solr core:", bold=True)
    add_para(doc,
        "The collection 'products' was created via the SolrCloud Collections "
        "API with numShards=2 and replicationFactor=2. Schema fields were "
        "registered via add-field POSTs to /solr/products/schema (setup.ps1). "
        "Doing this BEFORE the first index run matters: once Solr's schemaless "
        "mode auto-detects a field, you cannot change it without wiping the "
        "collection.")
    add_para(doc, "Field types selected:", bold=True)
    add_para(doc,
        "* text_general for title and description. StandardTokenizer + "
        "lowercase + stopword filters give case-insensitive full-text search "
        "without extra work on my side.\n"
        "* string for brand, category and subcategory. These are facet "
        "fields. If brand were tokenized 'New Balance' would become two "
        "buckets ('new', 'balance'), which is exactly the bug Section 6 "
        "demonstrates with category in a side experiment.\n"
        "* copyField from brand to brand_text (text_general). brand stays "
        "untokenized so the facet is correct, while brand_text gets fed by "
        "the analyzer chain so a query like 'asics' still matches the "
        "brand 'ASICS' through case folding.\n"
        "* pint, pfloat, boolean for the numeric and flag fields. Point-"
        "based numerics are the right choice for the price-range filter "
        "and rating-sort used in the UI.")

    # 4. Implementation steps
    add_heading(doc, "4. Implementation Steps")
    steps = [
        ("Install Solr", "Downloaded solr-9.6.1.tgz from archive.apache.org and extracted."),
        ("Start a SolrCloud cluster", ".\\solr-9.6.1\\bin\\solr.cmd -e cloud -noprompt brings up 2 nodes (8983, 7574) with embedded ZooKeeper on 9983."),
        ("Download dataset", "curl the raw Amazon-products sample from github.com/luminati-io/Amazon-dataset-samples into data/amazon_raw.csv (1,000 rows, 55 columns)."),
        ("Transform dataset", "python data/transform_amazon.py extracts the 11 fields used by this lab and drops 4 rows missing brand/price/rating. 996 valid rows out."),
        ("Create sharded collection", "POST /solr/admin/collections?action=CREATE&name=products&numShards=2&replicationFactor=2&collection.configName=_default — gives 4 cores spread across the two nodes."),
        ("Define schema", "POST add-field JSON for the 11 domain fields plus a copyField from brand to brand_text."),
        ("Index data", "POST products.csv to /solr/products/update?commit=true&header=true with Content-Type application/csv. 996 docs commit in under a second."),
        ("Verify shard distribution", "The 996 documents land roughly evenly across both shards via hash-based routing on the id field."),
        ("Run sample queries", "12 query types exercised in sample_queries.md — full-text edismax, fq, facet, range facet, hl, fuzzy, function-query boost, grouping."),
        ("Field-types experiment", "Created a separate 'products_bad' collection on a different configset with category declared as text_general instead of string. See Observations for the broken facet output."),
        ("Configure Suggester", "Posted SuggestComponent + /suggest_handler config to /solr/products/config (AnalyzingInfixLookupFactory over the title field)."),
        ("Build Flask UI", "app/app.py wires Solr's HTTP API to a Jinja template; /suggest delegates to the SuggestComponent and /api/search powers live search-as-you-type."),
        ("Test in browser", "Search, facets, price-range filter, sort dropdown, pagination, highlighting, autocomplete and live search verified manually with real queries (running, wireless, Skechers, ASICS)."),
    ]
    for i, (title_, body) in enumerate(steps, 1):
        add_para(doc, f"{i}. {title_}", bold=True)
        add_para(doc, "   " + body)

    add_para(doc, "The indexing command itself is a single REST call:")
    add_code(doc,
        'Invoke-RestMethod -Method Post `\n'
        '  -Uri "http://localhost:8983/solr/products/update?commit=true" `\n'
        '  -ContentType "application/csv; charset=utf-8" `\n'
        '  -InFile data/products.csv'
    )

    # 5. Screenshots
    add_heading(doc, "5. Screenshots")
    add_image_if_exists(doc, "01_solr_admin.png",  "Figure 1. Solr Admin UI dashboard showing the running SolrCloud instance.")
    add_image_if_exists(doc, "09_solrcloud_topology.png", "Figure 2. SolrCloud Cloud > Nodes view: two nodes (ports 8983 and 7574), each hosting two replicas of the sharded 'products' collection.")
    add_image_if_exists(doc, "11_query_tab.png",   "Figure 3. Solr Admin > Query tab on the products collection — the in-browser query builder used for ad-hoc testing during development.")
    add_image_if_exists(doc, "10_schema_tab.png",  "Figure 4. Schema tab for the 'brand' field on the products collection. Field-Type is StrField (string), Tokenized=NO. This keeps multi-word brands like 'New Balance' as a single facet bucket.")
    add_image_if_exists(doc, "13_field_experiment_bad_schema.png", "Figure 5. Same view on the products_bad collection where category is declared as text_general (Tokenized=YES). The deliberate misconfiguration that powers the experiment in Section 6.")
    add_image_if_exists(doc, "02_indexed_count.png", "Figure 6. q=*:* against the products collection returns numFound=996 (sum across both shards).")
    add_image_if_exists(doc, "03_facet_query.png",   "Figure 7. Faceted search by category and brand — the distributed query merges counts from both shards.")
    add_image_if_exists(doc, "04_highlight.png",     "Figure 8. Hit highlighting wraps matched terms with <mark> tags in the description field.")
    add_image_if_exists(doc, "12_field_experiment.png", "Figure 9. Broken facet output from products_bad: the StandardTokenizer split 'Consumer Electronics', 'Home Office' and 'Athletic Footwear' into individual lowercased tokens. Each multi-word category is now two unrelated buckets.")
    add_image_if_exists(doc, "05_web_ui.png",        "Figure 10. Flask web UI: a search for 'running' returns ASICS, WETIKE, Skechers and other real Amazon products, with highlighted hits and sidebar facets.")
    add_image_if_exists(doc, "06_autocomplete.png",  "Figure 11. Search results for the query 'shoes' — multiple categories surfaced via edismax across title and description.")
    add_image_if_exists(doc, "07_facet_ui.png",      "Figure 12. Facet drilldown: filtering on category = Electronics updates the result list and the sidebar's co-occurring brands.")
    add_image_if_exists(doc, "08_sort_ui.png",       "Figure 13. Sort by rating descending returns the highest-rated products first.")
    add_image_if_exists(doc, "14_live_search.png",   "Figure 14. Live search-as-you-type. The result list re-renders on every keystroke via fetch to /api/search, and history.replaceState keeps the URL in sync so any state remains shareable.")
    add_image_if_exists(doc, "15_failover_before.png", "Figure 15. Cluster topology before the failover test: both nodes (8983 and 7574) live, products collection's four replicas spread across them.")
    add_image_if_exists(doc, "16_failover_after.png", "Figure 16. Same view after running 'solr.cmd stop -p 7574'. Only one node is live; the products collection is still fully available because each shard has a replica on the surviving node.")
    add_image_if_exists(doc, "17_failover_query.png", "Figure 17. Flask web UI search for 'running' continues to return real Amazon results while node2 is down — confirming fault-tolerant distributed search.")
    add_image_if_exists(doc, "18_failover_json.png", "Figure 18. Raw Solr JSON response for the same query during the failover, captured directly from the surviving node1 endpoint.")
    add_image_if_exists(doc, "19_suggester.png",      "Figure 19. SuggestComponent response for suggest.q=harry — returns full title strings with the matched substring wrapped in <b>...</b> via AnalyzingInfixLookupFactory.")

    # 6. Observations
    add_heading(doc, "6. Observations and Analysis")
    add_para(doc, "Performance.", bold=True)
    add_para(doc,
        "Cold-cache queries on the 600-row index came back in 8 to 25 ms (the "
        "QTime field in the response header). Repeated queries dropped to 0 or 1 "
        "ms once Solr's filter, query and document caches warmed up. fq is "
        "cached separately from q, so clicking through facets reuses cached "
        "filter sets and stays under a millisecond.")
    add_para(doc, "Ranking.", bold=True)
    add_para(doc,
        "I used edismax with qf=title^3 brand_text^2 description. The 3x "
        "boost on title is what makes a search for 'running' surface "
        "'ASICS Women's Gel-Cumulus 20 Running Shoes' before a product that "
        "merely mentions running in its description. Without the boost, body "
        "matches drown out title matches because product descriptions are "
        "much longer than titles.")
    add_para(doc, "Faceting.", bold=True)
    add_para(doc,
        "facet.field on category and brand gave me the sidebar counts "
        "directly. facet.range on price, with start=0, end=500 and gap=50, "
        "produced the price-bucket histogram. facet.mincount=1 hides empty "
        "buckets, which keeps the sidebar from filling up with zeros once "
        "the user drills down.")
    add_para(doc, "Highlighting.", bold=True)
    add_para(doc,
        "hl=true with hl.simple.pre=<mark> tells Solr to wrap matched terms in "
        "the HTML <mark> tag inside the highlighted fragment. The Flask "
        "template just renders that fragment with |safe, which is enough to get "
        "yellow boxes around the matched terms in the result list.")
    add_para(doc, "Field-types experiment.", bold=True)
    add_para(doc,
        "To prove the schema choices in section 3 actually matter, I created "
        "a second collection 'products_bad' on a separate configset, "
        "declared its 'category' field as text_general instead of string, "
        "and indexed six rows covering three multi-word categories (Consumer "
        "Electronics x2, Home Office x2, Athletic Footwear x2).")
    add_para(doc,
        "products_bad facet output (Figure 9): "
        "'electronics' -> 2, 'consumer' -> 2, 'home' -> 2, 'office' -> 2, "
        "'athletic' -> 2, 'footwear' -> 2.")
    add_para(doc,
        "The StandardTokenizer split each multi-word value on whitespace "
        "and lowercased the tokens. 'Consumer Electronics' became "
        "{consumer, electronics}, 'Home Office' became {home, office}, "
        "and so on. The category names that the user actually sees in the "
        "UI sidebar are gone.")
    add_para(doc,
        "products (production) facet output for comparison: "
        "'Clothing, Shoes & Jewelry' -> 216, 'Home & Kitchen' -> 150, "
        "'Tools & Home Improvement' -> 112, ...")
    add_para(doc,
        "string fields skip the analyzer chain entirely. Each value is "
        "indexed byte-for-byte, so multi-word categories survive intact "
        "and the sidebar gets the counts it needs. Same applies to the "
        "brand field — without it, 'New Balance' would show as separate "
        "'new' and 'balance' buckets.")
    add_para(doc,
        "Side-finding while running this experiment: replacing the type on "
        "an already-indexed collection produced this Lucene error on the "
        "next write — 'cannot change field category from index "
        "options=DOCS to inconsistent index "
        "options=DOCS_AND_FREQS_AND_POSITIONS'. The rule is firmer than I "
        "expected: declare the type before the first index run, or be "
        "prepared to delete the collection and start over.")

    add_para(doc, "Fault tolerance (failover demo).", bold=True)
    add_para(doc,
        "I tested the cluster's tolerance to a node failure. Starting from "
        "both nodes up and 996 products indexed, I stopped node2 with "
        "'solr.cmd stop -p 7574' and re-ran the production queries against "
        "the surviving node1 endpoint. Cluster status now reported "
        "live_nodes = ['localhost:8983_solr'] (only one), but queries kept "
        "returning correct results: q=*:* still numFound=996, edismax "
        "q=running still surfaced the ASICS, Skechers and New Balance shoes "
        "at the top. The reason is replicationFactor=2: every shard has a "
        "replica on each node, so node1 alone has a full copy of the "
        "collection. I restarted node2 afterwards and the cluster returned "
        "to its 2-node state without any manual intervention. See Figures "
        "15-18.")

    add_para(doc, "Performance methodology and shard-overhead measurement.", bold=True)
    add_para(doc,
        "I wrote benchmark.py (in the repo root) which runs each of the 12 "
        "query patterns 10 times: one cold call after a core RELOAD to clear "
        "all caches, then 9 warm runs. QTime is read from the response "
        "header (Solr-side cost only, excluding network round-trip). The "
        "same 12 queries were run in two configurations against the same "
        "996-document collection: distributed (Solr fan-outs to both "
        "shards and merges) and single-shard (distrib=false, hits one "
        "shard's ~5,000 docs). Mean and p95 QTime over the 9 warm runs are "
        "reported.")
    add_para(doc, "Selected results (mean QTime over 9 warm runs, in ms):")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Query"
    hdr[1].text = "distributed (2 shards)"
    hdr[2].text = "single-shard"
    hdr[3].text = "ratio"
    rows = [
        ("Q1 match-all",         "24.3", "0.9", "27x"),
        ("Q2 edismax full-text", "32.4", "1.7", "19x"),
        ("Q3 fq + price range",  "11.7", "2.0",  "6x"),
        ("Q4 multi-field facets","18.3", "1.8", "10x"),
        ("Q5 sort rating+reviews","10.9","0.6", "20x"),
        ("Q6 highlighting",      "33.4","16.6",  "2x"),
        ("Q9 fuzzy",             "14.9", "2.4",  "6x"),
        ("Q11 paging start=100", "13.3", "0.3", "40x"),
    ]
    for q, d_, s_, r_ in rows:
        rc = table.add_row().cells
        rc[0].text = q
        rc[1].text = d_
        rc[2].text = s_
        rc[3].text = r_
    add_para(doc, "Interpretation.", bold=True)
    add_para(doc,
        "Distributed wins zero of the twelve queries. It is 1.5x to 40x "
        "slower than the single-shard equivalent across the board. The "
        "reason is boring once you look at the single-shard column: each "
        "shard's work on 498 documents is already sub-millisecond for most "
        "queries. There is nothing to parallelize. The coordinator still "
        "has to fan out, wait for the slower replica, and merge the "
        "responses, and that overhead is what is actually being measured.")
    add_para(doc,
        "Sharding starts paying off when per-shard work is expensive "
        "enough to make coordinator cost look small. For Lucene that "
        "usually means millions of documents, not one thousand. The "
        "cluster was worth setting up because the prerequisite asked for "
        "it and because it lets me demonstrate failover, but I would not "
        "shard a 1,000-row catalog in production. Per-query numbers "
        "including p95 are in report/benchmark_results.txt.")

    add_para(doc, "Suggester component (autocomplete).", bold=True)
    add_para(doc,
        "After the empirical benchmark I replaced the original wildcard-"
        "based autocomplete with Solr's proper SuggestComponent. I posted "
        "this configuration to /solr/products/config:")
    add_code(doc,
        '{"add-searchcomponent":{\n'
        '   "name":"suggest_component",\n'
        '   "class":"solr.SuggestComponent",\n'
        '   "suggester":{\n'
        '     "name":"titleSuggester",\n'
        '     "lookupImpl":"AnalyzingInfixLookupFactory",\n'
        '     "dictionaryImpl":"DocumentDictionaryFactory",\n'
        '     "field":"title",\n'
        '     "suggestAnalyzerFieldType":"text_general",\n'
        '     "buildOnCommit":"true"\n'
        '   }},\n'
        ' "add-requesthandler":{\n'
        '   "name":"/suggest_handler",\n'
        '   "class":"solr.SearchHandler",\n'
        '   "defaults":{"suggest":"true","suggest.count":"10",\n'
        '              "suggest.dictionary":"titleSuggester"},\n'
        '   "components":["suggest_component"]}\n'
        '}')
    add_para(doc,
        "AnalyzingInfixLookupFactory does prefix AND infix matching against "
        "the analyzed title field, so 'harry' matches 'The Lincoln Lawyer "
        "the title field, so 'running' surfaces 'ASICS Women's Gel-Cumulus "
        "20 Running Shoes' alongside 'WETIKE Mesh Slip On Lightweight "
        "Running Sneakers'. The Flask /suggest endpoint now "
        "calls the Suggester first and falls back to the wildcard query if "
        "the handler isn't configured (older deployments). See Figure 19.")
    add_para(doc, "Side-finding while configuring this:", bold=True)
    add_para(doc,
        "The first time I posted this config I sent buildOnCommit: true as a "
        "JSON boolean and got a 500 back. The Solr server log had the real "
        "story: 'class java.lang.Boolean cannot be cast to class "
        "java.lang.String' inside SuggestComponent.inform(). Quoting it as "
        "\"true\" worked immediately. The schema and config APIs accept JSON "
        "but pass several values through string parsers internally, so any "
        "unquoted boolean or number can hit this.")

    add_para(doc, "Live search-as-you-type.", bold=True)
    add_para(doc,
        "The Flask UI exposes a JSON endpoint at /api/search that returns the "
        "same result set as the rendered page. The frontend debounces input "
        "events on the search box (180 ms) and re-renders the result list on "
        "every keystroke, without reloading the page. history.replaceState "
        "keeps the URL in sync so any state (query, facets, sort) is still "
        "shareable. The traditional form submit also still works as a "
        "no-JavaScript fallback.")

    add_para(doc, "What I would change.", bold=True)
    add_para(doc,
        "If I were doing this on a real catalog I would add a copyField from "
        "title and description into a single all-text field and qf against that, "
        "instead of listing each searchable field. I would also turn on the "
        "Suggester component for autocomplete instead of running a wildcard "
        "title query, which is what /suggest does today.")

    # 7. Challenges and solutions
    add_heading(doc, "7. Challenges Faced and Solutions")
    challenges = [
        ("Schema-less mode mis-typed numeric fields as strings",
         "Explicitly POSTed add-field requests with pint/pfloat types BEFORE indexing. "
         "Once a field is auto-detected as string, it cannot be changed without deleting and reindexing."),
        ("Facet on a tokenized field returned split tokens (Science / Fiction)",
         "Switched category and brand to type=string so the analyzer chain doesn't tokenize them. Added brand_text (text_general, populated via copyField) for full-text search."),
        ("Solr's default request-handler returned only 10 docs",
         "Added rows= and start= parameters and built pagination controls in the Flask template."),
        ("Special characters in queries (':' '/' '\"') broke parsing",
         "Used edismax instead of the default lucene parser — edismax silently escapes user input "
         "and falls back gracefully on malformed queries."),
        ("CSV with embedded commas in description failed to import",
         "Used header=true and the standard Solr CSV update handler, which handles quoted fields "
         "correctly per RFC 4180."),
    ]
    for c, s in challenges:
        add_para(doc, "Challenge: " + c, bold=True)
        add_para(doc, "Solution: " + s)
        add_para(doc, "")

    # 8. Conclusion
    add_heading(doc, "8. Conclusion")
    add_para(doc,
        "By the end I had a 2-node SolrCloud cluster holding 996 real "
        "Amazon products across two shards, with replicas mirrored so a "
        "node failure does not take queries down. Twelve query patterns "
        "work against the distributed collection, all of them against "
        "real fields read from the source CSV. The Flask UI on top covers "
        "search, facets, price range, sort, pagination, highlighting, "
        "real autocomplete via SuggestComponent, and search-as-you-type.")
    add_para(doc,
        "Two findings I want to keep from the empirical work. First, "
        "schema choices are load-bearing in a way I did not appreciate "
        "before. Declaring brand or category as text_general silently "
        "splits multi-word values like 'New Balance' or 'Home & Kitchen' "
        "across separate facet buckets. The fix is string + a copyField "
        "to a tokenized sibling for full-text search. Second, sharding at "
        "this corpus size is a net loss. The distributed version of every "
        "query was 1.5x to 40x slower than the single-shard equivalent "
        "because the coordinator overhead dominates when per-shard work "
        "is already sub-millisecond. The cluster is the right shape for "
        "the topology and failover demos, but I would not shard a "
        "1,000-row catalog in production.")
    add_para(doc,
        "Two things bit me and are worth remembering for next time. First: "
        "declare numeric and string field types BEFORE the first index run, "
        "because schemaless mode silently locks them in as strings the moment "
        "you commit a row. Second: keep facet fields untokenized, otherwise "
        "'Science Fiction' becomes two buckets in the sidebar.")
    add_para(doc,
        "All the source, the dataset, the PowerShell setup scripts and the "
        "screenshots used in this report are in the linked repo.")

    add_para(doc, "")
    add_para(doc, "GitHub repository: https://github.com/huzvert/solr-books-lab", bold=True)
    add_para(doc, "(Repo retains its original 'solr-books-lab' name from the early commits; "
                  "the project pivoted to a real Amazon-products dataset in later commits to "
                  "remove all synthetic fields. All current code, docs and screenshots use "
                  "the products schema.)", italic=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
